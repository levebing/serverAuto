from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file, make_response, session
from flask_cors import CORS
from database import add_server, get_all_servers, get_server_by_id, update_server, delete_server, add_inspection_record, get_inspection_records, get_inspection_record_by_id, get_all_groups, add_group, delete_group, update_group, search_servers, add_scheduled_task, get_scheduled_task_by_server_id, get_all_scheduled_tasks, update_scheduled_task, delete_scheduled_task, add_report, get_all_reports, delete_report, update_group_sort_order, get_servers_count, get_inspection_records_count, get_scheduled_tasks_count, get_reports_count, add_inspection_item, get_all_inspection_items, get_inspection_item_by_id, update_inspection_item, delete_inspection_item, toggle_inspection_item, get_inspection_items_count, get_user_by_username, update_user_password, add_upload_record, get_upload_record_by_id, get_all_upload_records, get_upload_records_count, delete_upload_record, get_upload_records_by_file_path, add_fault, get_all_faults, get_faults_count, get_fault_by_id, update_fault, delete_fault, review_fault, get_alert_config, save_alert_config, add_alert_log, get_alert_logs, get_alert_logs_count
from encryption import decrypt_password
from inspection import ServerInspector
from windows_inspection import WindowsServerInspector
import io
import os
import uuid
import hashlib
from docx import Document
from docx.shared import Pt, RGBColor
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import requests
import config

app = Flask(__name__)
CORS(app)
app.secret_key = config.SECRET_KEY

def login_required(f):
    """登录装饰器"""
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    decorated_function.__doc__ = f.__doc__
    return decorated_function

def generate_docx_report(server_name, ip, result):
    doc = Document()
    
    title = doc.add_heading('服务器巡检报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = '服务器名称'
    info_table.cell(0, 1).text = server_name
    info_table.cell(1, 0).text = '服务器IP'
    info_table.cell(1, 1).text = ip
    
    doc.add_paragraph()
    
    doc.add_heading('一、巡检时间', level=2)
    doc.add_paragraph(f"巡检时间: {result['inspection_time']}")
    doc.add_paragraph(f"系统时间: {result['system_time']}")
    
    doc.add_heading('二、操作系统信息', level=2)
    os_info = result['os_version'] or '无法获取操作系统版本'
    doc.add_paragraph(os_info)
    
    doc.add_heading('三、CPU使用率', level=2)
    cpu_usage = result['cpu_usage'] or '无法获取'
    doc.add_paragraph(f"CPU使用率: {cpu_usage}%")
    
    doc.add_heading('四、内存使用情况', level=2)
    memory_info = eval(result['memory_usage']) if result['memory_usage'] else {}
    if memory_info:
        memory_table = doc.add_table(rows=4, cols=2)
        memory_table.style = 'Table Grid'
        memory_table.cell(0, 0).text = '总计'
        memory_table.cell(0, 1).text = memory_info.get('total', '未知')
        memory_table.cell(1, 0).text = '已用'
        memory_table.cell(1, 1).text = memory_info.get('used', '未知')
        memory_table.cell(2, 0).text = '可用'
        memory_table.cell(2, 1).text = memory_info.get('available', '未知')
        memory_table.cell(3, 0).text = '使用率'
        usage_percent = memory_info.get('usage_percent', '未知')
        memory_table.cell(3, 1).text = f"{usage_percent}%"
        if isinstance(usage_percent, int) and usage_percent >= 80:
            memory_table.cell(3, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
    else:
        doc.add_paragraph('无法获取内存信息')
    
    doc.add_heading('五、磁盘使用情况', level=2)
    disk_info = eval(result['disk_usage']) if result['disk_usage'] else []
    if disk_info:
        disk_table = doc.add_table(rows=len(disk_info) + 1, cols=5)
        disk_table.style = 'Table Grid'
        headers = ['挂载点', '文件系统', '总计', '已用', '使用率']
        for i, header in enumerate(headers):
            disk_table.cell(0, i).text = header
            disk_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        
        for i, disk in enumerate(disk_info, 1):
            disk_table.cell(i, 0).text = disk.get('mount_point', '')
            disk_table.cell(i, 1).text = disk.get('filesystem', '')
            disk_table.cell(i, 2).text = disk.get('size', '')
            disk_table.cell(i, 3).text = disk.get('used', '')
            usage = disk.get('usage_percent', 0)
            disk_table.cell(i, 4).text = f"{usage}%"
            if usage >= 80:
                disk_table.cell(i, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
    else:
        doc.add_paragraph('无法获取磁盘信息')
    
    doc.add_heading('六、告警内容', level=2)
    alert_content = result['alert_content']
    alert_paragraph = doc.add_paragraph(alert_content)
    if alert_content != '无告警':
        alert_paragraph.runs[0].font.color.rgb = RGBColor(220, 53, 69)
        alert_paragraph.runs[0].font.bold = True
    else:
        alert_paragraph.runs[0].font.color.rgb = RGBColor(34, 197, 94)
    
    footer = doc.sections[0].footer
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = "服务器自动巡检工具生成"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        user = get_user_by_username(username)
        if user:
            # 用户存在，验证密码（解密后比较）
            stored_password = user[2]  # 数据库中加密的密码
            decrypted_password = decrypt_password(stored_password)
            if password == decrypted_password:
                session['username'] = username
                return jsonify({'success': True})
        
        return jsonify({'success': False, 'message': '用户名或密码错误'})
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

@app.route('/change_password', methods=['GET', 'POST'])
@login_required
def change_password():
    if request.method == 'POST':
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        
        if not old_password or not new_password or not confirm_password:
            return render_template('change_password.html', message='密码不能为空', success=False)
        
        if new_password != confirm_password:
            return render_template('change_password.html', message='两次输入的密码不一致', success=False)
        
        username = session.get('username')
        user = get_user_by_username(username)
        
        if not user:
            return render_template('change_password.html', message='用户不存在', success=False)
        
        # 验证原密码
        stored_password = user[2]
        decrypted_password = decrypt_password(stored_password)
        
        if old_password != decrypted_password:
            return render_template('change_password.html', message='原密码错误', success=False)
        
        # 更新密码
        user_id = user[0]
        update_user_password(user_id, new_password)
        
        return redirect(url_for('index'))
    
    return render_template('change_password.html')

@app.route('/')
@login_required
def index():
    group_id = request.args.get('group_id', 'all')
    ip = request.args.get('ip', '')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    servers = []
    total = 0
    if ip:
        servers = search_servers(ip)
        total = len(servers)
    else:
        servers = get_all_servers(group_id, page, per_page)
        total = get_servers_count(group_id)
    total_pages = (total + per_page - 1) // per_page
    groups = get_all_groups()
    return render_template('index.html', servers=servers, groups=groups, selected_group=group_id, search_ip=ip, page=page, per_page=per_page, total=total, total_pages=total_pages)

@app.route('/records')
@login_required
def records():
    group_id = request.args.get('group_id', 'all')
    ip = request.args.get('ip', '')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    records = get_inspection_records(group_id=group_id, page=page, per_page=per_page)
    total = get_inspection_records_count(group_id=group_id)
    total_pages = (total + per_page - 1) // per_page
    groups = get_all_groups()
    return render_template('records.html', records=records, groups=groups, selected_group=group_id, search_ip=ip, page=page, per_page=per_page, total=total, total_pages=total_pages)

@app.route('/add', methods=['GET', 'POST'])
@login_required
def add():
    groups = get_all_groups()
    if request.method == 'POST':
        name = request.form['name']
        ip = request.form['ip']
        port = int(request.form['port'])
        username = request.form['username']
        group_id = int(request.form.get('group_id', 1))
        remark = request.form.get('remark', '')
        private_key_content = None
        password = None
        
        if 'private_key' in request.files:
            key_file = request.files['private_key']
            if key_file.filename != '':
                private_key_content = key_file.read()
        
        if not private_key_content:
            password = request.form.get('password')
        
        os_type = request.form.get('os_type', 'linux')
        
        add_server(name, ip, port, username, group_id, remark, private_key_content, password, os_type)
        return redirect(url_for('index'))
    
    return render_template('add.html', groups=groups)

@app.route('/edit/<int:server_id>', methods=['GET', 'POST'])
@login_required
def edit(server_id):
    server = get_server_by_id(server_id)
    groups = get_all_groups()
    if not server:
        return "服务器不存在", 404
    
    if request.method == 'POST':
        name = request.form['name']
        ip = request.form['ip']
        port = int(request.form['port'])
        username = request.form['username']
        group_id = int(request.form.get('group_id', 1))
        remark = request.form.get('remark', '')
        private_key_content = server[5]
        password = server[6]  # 密码字段索引为6（解密后的）
        
        if 'private_key' in request.files:
            key_file = request.files['private_key']
            if key_file.filename != '':
                private_key_content = key_file.read()
        
        if not private_key_content:
            password = request.form.get('password')
        
        os_type = request.form.get('os_type', 'linux')
        
        update_server(server_id, name, ip, port, username, group_id, remark, private_key_content, password, os_type)
        return redirect(url_for('index'))
    
    return render_template('edit.html', server=server, groups=groups)

@app.route('/delete/<int:server_id>')
@login_required
def delete(server_id):
    delete_server(server_id)
    return redirect(url_for('index'))

@app.route('/inspect/<int:server_id>')
@login_required
def inspect(server_id):
    server = get_server_by_id(server_id)
    if not server:
        return "服务器不存在", 404
    
    # 根据操作系统类型选择巡检方式
    os_type = server[8] if len(server) > 8 else 'linux'
    
    if os_type == 'windows':
        inspector = WindowsServerInspector(
            ip=server[2],
            port=server[3],
            username=server[4],
            password=server[6]
        )
    else:
        inspector = ServerInspector(
            ip=server[2],
            port=server[3],
            username=server[4],
            private_key_content=server[5],
            password=server[6]
        )
    
    result, error = inspector.inspect()
    
    if error:
        groups = get_all_groups()
        return render_template('inspect.html', server=server, error=error, groups=groups)
    
    inspection_result = generate_inspection_result(result)
    
    add_inspection_record(
        server_id=server_id,
        disk_usage=result['disk_usage'],
        memory_usage=result['memory_usage'],
        cpu_usage=result['cpu_usage'],
        system_time=result['system_time'],
        os_version=result['os_version'],
        alert_content=result['alert_content'],
        report_content=str(result),
        inspection_result=inspection_result,
        inspection_time=result['inspection_time']
    )
    
    # 触发短信告警（如果有告警内容）
    if result.get('alert_content') and result['alert_content'] != '无告警':
        try:
            send_inspection_alert_sms(server, result)
        except Exception as e:
            print(f"短信告警发送失败: {e}")
    
    groups = get_all_groups()
    return render_template('inspect.html', server=server, result=result, groups=groups)

def generate_inspection_result(result):
    """生成巡检结果摘要"""
    items = []
    
    cpu = result.get('cpu_usage')
    if cpu:
        items.append(f"CPU {cpu}%")
    
    memory = result.get('memory_usage')
    if memory:
        try:
            mem_info = eval(memory)
            if mem_info.get('usage_percent'):
                items.append(f"内存 {mem_info['usage_percent']}%")
        except:
            pass
    
    disk = result.get('disk_usage')
    if disk:
        try:
            disk_info = eval(disk)
            for d in disk_info:
                if d.get('usage_percent'):
                    items.append(f"{d['mount_point']} {d['usage_percent']}%")
        except:
            pass
    
    # 添加自定义巡检项的结果
    custom_results = result.get('custom_results', {})
    predefined_items = ['CPU使用率', '内存使用情况', '磁盘使用情况', '系统时间', '操作系统版本']
    for name, output in custom_results.items():
        if name not in predefined_items and output:
            # 对于自定义巡检项，显示实际的输出内容
            output_str = str(output).strip()
            if output_str:
                # 截取前50个字符作为摘要，避免过长
                if len(output_str) > 50:
                    output_str = output_str[:50] + '...'
                # 替换换行符为空格
                output_str = output_str.replace('\n', ' ').replace('\r', '')
                items.append(f"{name}: {output_str}")
            else:
                items.append(f"{name}: 无输出")
    
    alert = result.get('alert_content')
    if alert and alert != '无告警':
        status = '⚠️ 告警'
    else:
        status = '✅ 正常'
    
    return f"{status} | {', '.join(items)}"

@app.route('/api/inspect_async/<int:server_id>', methods=['POST'])
@login_required
def inspect_async(server_id):
    server = get_server_by_id(server_id)
    if not server:
        return jsonify({'success': False, 'message': '服务器不存在'})
    
    # 根据操作系统类型选择巡检方式
    os_type = server[8] if len(server) > 8 else 'linux'
    
    if os_type == 'windows':
        inspector = WindowsServerInspector(
            ip=server[2],
            port=server[3],
            username=server[4],
            password=server[6]
        )
    else:
        inspector = ServerInspector(
            ip=server[2],
            port=server[3],
            username=server[4],
            private_key_content=server[5],
            password=server[6]
        )
    
    result, error = inspector.inspect()
    
    if error:
        return jsonify({'success': False, 'message': error})
    
    inspection_result = generate_inspection_result(result)
    
    add_inspection_record(
        server_id=server_id,
        disk_usage=result['disk_usage'],
        memory_usage=result['memory_usage'],
        cpu_usage=result['cpu_usage'],
        system_time=result['system_time'],
        os_version=result['os_version'],
        alert_content=result['alert_content'],
        report_content=str(result),
        inspection_result=inspection_result,
        inspection_time=result['inspection_time']
    )
    
    return jsonify({
        'success': True,
        'server_name': server[1],
        'server_ip': server[2],
        'inspection_result': inspection_result,
        'inspection_time': result['inspection_time']
    })

@app.route('/download')
@login_required
def download_file():
    path = request.args.get('path')
    if not path:
        return "参数错误", 400
    
    safe_path = os.path.normpath(path)
    if safe_path.startswith('..') or safe_path.startswith('/'):
        return "路径非法", 403
    
    file_path = os.path.join(config.UPLOAD_LOCAL_PATH, safe_path)
    if os.path.exists(file_path):
        return send_from_directory(os.path.dirname(file_path), os.path.basename(file_path))
    return "文件不存在", 404

@app.route('/download_report/<int:record_id>')
@login_required
def download_report(record_id):
    record = get_inspection_record_by_id(record_id)
    if not record:
        return "记录不存在", 404
    
    server = get_server_by_id(record[1])
    server_name = server[1] if server else "unknown"
    server_ip = server[2] if server else "unknown"
    
    try:
        result = eval(record[9]) if record[9] else {}
    except:
        result = {}
    
    result['inspection_time'] = record[7]
    result['system_time'] = record[5]
    result['os_version'] = record[6]
    result['cpu_usage'] = record[4]
    result['memory_usage'] = record[3]
    result['disk_usage'] = record[2]
    result['alert_content'] = record[8] if record[8] else '无告警'
    
    doc = generate_docx_report(server_name, server_ip, result)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename = f"巡检报告_{server_name}_{record[8]}.docx"
    
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/delete_record/<int:record_id>')
@login_required
def delete_record(record_id):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('UPDATE inspection_records SET is_deleted = 1 WHERE id = ?', (record_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('records'))

@app.route('/api/get_record_detail/<int:record_id>')
@login_required
def get_record_detail(record_id):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT ir.*, s.name as server_name, s.ip as server_ip, g.name as group_name
        FROM inspection_records ir
        LEFT JOIN servers s ON ir.server_id = s.id
        LEFT JOIN groups g ON s.group_id = g.id
        WHERE ir.id = ? AND ir.is_deleted = 0
    ''', (record_id,))
    record = cursor.fetchone()
    conn.close()
    
    if not record:
        return jsonify({'success': False, 'message': '记录不存在'})
    
    return jsonify({
        'success': True,
        'record': {
            'server_name': record[12],
            'server_ip': record[13],
            'group_name': record[14],
            'cpu_usage': record[4],
            'memory_usage': record[3],
            'disk_usage': record[2],
            'system_time': record[5],
            'os_version': record[6],
            'alert_content': record[8],
            'inspection_result': record[10],
            'inspection_time': record[7]
        }
    })

@app.route('/api/add_group', methods=['POST'])
@login_required
def api_add_group():
    name = request.json.get('name')
    sort_order = request.json.get('sort_order', 0)
    parent_id = request.json.get('parent_id')
    
    if not name:
        return jsonify({'success': False, 'message': '分组名称不能为空'})
    
    if parent_id == '' or parent_id == '0' or parent_id is None:
        parent_id = None
    
    group_id = add_group(name, sort_order, parent_id)
    if group_id:
        return jsonify({'success': True, 'group_id': group_id, 'name': name})
    else:
        return jsonify({'success': False, 'message': '分组名称已存在'})

@app.route('/api/delete_group/<int:group_id>')
@login_required
def api_delete_group(group_id):
    if group_id == 1:
        return jsonify({'success': False, 'message': '不能删除默认分组'})
    delete_group(group_id)
    return jsonify({'success': True})

@app.route('/api/update_group', methods=['POST'])
@login_required
def api_update_group():
    group_id = request.json.get('group_id')
    name = request.json.get('name')
    sort_order = request.json.get('sort_order')
    parent_id = request.json.get('parent_id')
    
    if not group_id:
        return jsonify({'success': False, 'message': '分组ID不能为空'})
    if int(group_id) == 1:
        return jsonify({'success': False, 'message': '不能修改默认分组'})
    
    if parent_id == '' or parent_id == '0' or parent_id is None:
        parent_id = None
    
    result = update_group(group_id, name, sort_order, parent_id)
    if result is None:
        return jsonify({'success': False, 'message': '分组名称已存在'})
    elif result:
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'message': '分组不存在'})

@app.route('/api/update_group_sort', methods=['POST'])
@login_required
def api_update_group_sort():
    group_id = request.json.get('group_id')
    sort_order = request.json.get('sort_order')
    if not group_id or sort_order is None:
        return jsonify({'success': False, 'message': '参数不能为空'})
    result = update_group_sort_order(group_id, sort_order)
    if result:
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'message': '分组不存在'})

@app.route('/api/search_groups', methods=['GET'])
@login_required
def api_search_groups():
    keyword = request.args.get('keyword', '').strip()
    groups_list = get_all_groups()
    
    if keyword:
        groups_list = [g for g in groups_list if keyword.lower() in g[1].lower()]
    
    return jsonify({'success': True, 'data': groups_list})

@app.route('/groups')
@login_required
def groups():
    groups = get_all_groups()
    return render_template('groups.html', groups=groups)

@app.route('/tasks')
@login_required
def tasks():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    tasks = get_all_scheduled_tasks(page, per_page)
    total = get_scheduled_tasks_count()
    total_pages = (total + per_page - 1) // per_page
    servers = get_all_servers()
    groups = get_all_groups()
    return render_template('tasks.html', tasks=tasks, servers=servers, groups=groups, page=page, per_page=per_page, total=total, total_pages=total_pages)

@app.route('/api/add_task', methods=['POST'])
@login_required
def api_add_task():
    server_id = request.json.get('server_id')
    cron_expression = request.json.get('cron_expression')
    is_enabled = request.json.get('is_enabled', 1)
    
    if not server_id or not cron_expression:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    task_id = add_scheduled_task(server_id, cron_expression, is_enabled)
    if task_id:
        return jsonify({'success': True, 'task_id': task_id})
    else:
        return jsonify({'success': False, 'message': '添加任务失败'})

@app.route('/api/update_task', methods=['POST'])
@login_required
def api_update_task():
    task_id = request.json.get('task_id')
    cron_expression = request.json.get('cron_expression')
    is_enabled = request.json.get('is_enabled')
    
    if not task_id or not cron_expression:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    result = update_scheduled_task(task_id, cron_expression, is_enabled)
    if result:
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'message': '更新任务失败'})

@app.route('/api/delete_task/<int:server_id>')
@login_required
def api_delete_task(server_id):
    delete_scheduled_task(server_id)
    return jsonify({'success': True})

@app.route('/api/delete_report/<int:report_id>')
@login_required
def api_delete_report(report_id):
    result = delete_report(report_id)
    return jsonify({'success': result})

# 巡检项管理路由
@app.route('/inspection_items')
@login_required
def inspection_items():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    os_type = request.args.get('os_type')
    
    items = get_all_inspection_items(os_type=os_type, page=page, per_page=per_page)
    total = get_inspection_items_count(os_type=os_type)
    total_pages = (total + per_page - 1) // per_page
    
    return render_template('inspection_items.html', 
                         inspection_items=items, 
                         page=page, 
                         per_page=per_page, 
                         total=total, 
                         total_pages=total_pages,
                         os_type=os_type)

@app.route('/inspection_item', methods=['GET', 'POST'])
@login_required
def inspection_item():
    item_id = request.args.get('id')
    item = None
    
    if item_id:
        item = get_inspection_item_by_id(int(item_id))
    
    if request.method == 'POST':
        item_id = request.form.get('item_id')
        name = request.form.get('item_name')
        command = request.form.get('item_command')
        description = request.form.get('item_description', '')
        os_type = request.form.get('item_os_type', 'linux')
        
        if not name or not command:
            return render_template('add_inspection_item.html', item=item, message='名称和命令不能为空', success=False)
        
        if item_id:
            # 更新巡检项
            result = update_inspection_item(int(item_id), name, command, description, os_type)
            if result:
                return redirect(url_for('inspection_items'))
            else:
                return render_template('add_inspection_item.html', item=item, message='更新失败', success=False)
        else:
            # 添加巡检项
            result = add_inspection_item(name, command, description, os_type)
            if result:
                return redirect(url_for('inspection_items'))
            else:
                return render_template('add_inspection_item.html', item=None, message='添加失败', success=False)
    
    return render_template('add_inspection_item.html', item=item)

@app.route('/api/add_inspection_item', methods=['POST'])
@login_required
def api_add_inspection_item():
    name = request.json.get('name')
    command = request.json.get('command')
    description = request.json.get('description', '')
    
    if not name or not command:
        return jsonify({'success': False, 'message': '名称和命令不能为空'})
    
    item_id = add_inspection_item(name, command, description)
    if item_id:
        return jsonify({'success': True, 'item_id': item_id})
    else:
        return jsonify({'success': False, 'message': '添加失败'})

@app.route('/api/update_inspection_item', methods=['POST'])
@login_required
def api_update_inspection_item():
    item_id = request.json.get('id')
    name = request.json.get('name')
    command = request.json.get('command')
    description = request.json.get('description', '')
    
    if not item_id or not name or not command:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    result = update_inspection_item(item_id, name, command, description)
    if result:
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'message': '更新失败'})

@app.route('/api/delete_inspection_item/<int:item_id>')
@login_required
def api_delete_inspection_item(item_id):
    result = delete_inspection_item(item_id)
    return jsonify({'success': result})

@app.route('/api/toggle_inspection_item', methods=['POST'])
@login_required
def api_toggle_inspection_item():
    item_id = request.json.get('id')
    is_enabled = request.json.get('is_enabled')
    
    if item_id is None or is_enabled is None:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    result = toggle_inspection_item(item_id, is_enabled)
    if result:
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'message': '操作失败'})

@app.route('/upload/<path:filename>')
def serve_upload(filename):
    file_path = os.path.join(config.UPLOAD_LOCAL_PATH, filename)
    if os.path.exists(file_path):
        return send_file(file_path)
    return jsonify({'success': False, 'message': '文件不存在'}), 404

@app.route('/api/upload', methods=['POST'])
def api_upload():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '文件名为空'}), 400
    
    try:
        original_name = file.filename
        ext = os.path.splitext(original_name)[1]
        timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        file_name = f"{timestamp}_{unique_id}{ext}"
        
        year_month = datetime.datetime.now().strftime('%Y-%m')
        year_month_path = os.path.join(config.UPLOAD_LOCAL_PATH, year_month)
        os.makedirs(year_month_path, exist_ok=True)
        
        file_path = os.path.join(year_month_path, file_name)
        file.save(file_path)
        
        file_size = os.path.getsize(file_path)
        relative_path = f"upload/{year_month}/{file_name}"
        full_url = f"{config.UPLOAD_LOCAL_BASE_URL}/{year_month}/{file_name}"
        
        record_id = add_upload_record(file_name, original_name, relative_path, file_size, 'local')
        
        return jsonify({
            'success': True,
            'code': 200,
            'message': '上传成功',
            'data': [full_url],
            'record_id': record_id,
            'file_name': file_name,
            'original_name': original_name,
            'file_size': file_size
        })
    except Exception as e:
        return jsonify({'success': False, 'code': 500, 'message': str(e)}), 500

@app.route('/api/upload_records', methods=['GET'])
@login_required
def api_upload_records():
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    
    records = get_all_upload_records(page, per_page)
    total = get_upload_records_count()
    total_pages = (total + per_page - 1) // per_page
    
    records_list = []
    for record in records:
        records_list.append({
            'id': record[0],
            'file_name': record[1],
            'original_name': record[2],
            'file_path': record[3],
            'file_size': record[4],
            'storage_type': record[5],
            'storage_time': record[6].isoformat() if record[6] else None,
            'created_at': record[8].isoformat() if record[8] else None
        })
    
    return jsonify({
        'success': True,
        'data': records_list,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': total_pages
    })

@app.route('/api/delete_upload_record/<int:record_id>', methods=['POST'])
@login_required
def api_delete_upload_record(record_id):
    record = get_upload_record_by_id(record_id)
    if not record:
        return jsonify({'success': False, 'message': '记录不存在'})
    
    file_path = os.path.join(config.UPLOAD_LOCAL_PATH, record[3])
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            pass
    
    result = delete_upload_record(record_id)
    if result:
        return jsonify({'success': True, 'message': '删除成功'})
    else:
        return jsonify({'success': False, 'message': '删除失败'})

@app.route('/reports')
@login_required
def reports():
    groups = get_all_groups()
    
    # 获取筛选参数
    filter_type = request.args.get('report_type')
    filter_group = request.args.get('group_id')
    filter_date = request.args.get('date')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    # 获取报告列表
    reports = get_all_reports(filter_type, filter_group, filter_date, page, per_page)
    total = get_reports_count(filter_type, filter_group, filter_date)
    total_pages = (total + per_page - 1) // per_page
    
    if config.UPLOAD_TYPE == 'local':
        base_url = config.UPLOAD_LOCAL_BASE_URL
    else:
        base_url = config.REPORT_STORAGE['base_url']
    
    return render_template('reports.html', groups=groups, reports=reports, 
                           filter_type=filter_type, filter_group=filter_group, filter_date=filter_date, 
                           page=page, per_page=per_page, total=total, total_pages=total_pages, 
                           base_url=base_url)

@app.route('/api/generate_report', methods=['POST'])
def api_generate_report():
    report_type = request.json.get('report_type')
    group_id = request.json.get('group_id')
    start_date_str = request.json.get('start_date')
    
    if not report_type or not group_id:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    # 根据报告类型和分组生成报告
    records = get_inspection_records(group_id=group_id)
    
    # 生成报告内容
    report_content = generate_report_content(records, report_type, start_date_str)
    
    # 生成文件名
    import datetime
    now = datetime.datetime.now()
    filename = f"{report_type}_report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    
    # 获取分组名称
    group_name = '所有分组' if group_id == 'all' else None
    if group_id != 'all':
        groups = get_all_groups()
        for group in groups:
            if str(group[0]) == group_id:
                group_name = group[1]
                break
    
    # 生成报告名称
    report_name = report_content['title']
    
    # 生成Word文档
    doc = generate_report_docx(report_content, report_type)
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # 上传文件
    file_path = None
    try:
        output.seek(0)
        
        if config.UPLOAD_TYPE == 'local':
            ext = os.path.splitext(filename)[1]
            timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            stored_filename = f"{timestamp}_{unique_id}{ext}"
            
            year_month = datetime.datetime.now().strftime('%Y-%m')
            year_month_path = os.path.join(config.UPLOAD_LOCAL_PATH, year_month)
            os.makedirs(year_month_path, exist_ok=True)
            
            file_full_path = os.path.join(year_month_path, stored_filename)
            
            with open(file_full_path, 'wb') as f:
                f.write(output.getvalue())
            
            relative_path = f"upload/{year_month}/{stored_filename}"
            file_path = f"{config.UPLOAD_LOCAL_BASE_URL}/{year_month}/{stored_filename}"
            print(f"文件已保存到本地: {file_path}")
            
            add_upload_record(stored_filename, filename, relative_path, len(output.getvalue()), 'local')
            db_file_path = relative_path
        else:
            upload_url = config.FILE_UPLOAD_SERVICE['url']
            print(f"开始上传文件到: {upload_url}")
            print(f"文件名: {filename}")
            
            files = {'files': (filename, output, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(upload_url, files=files, timeout=config.FILE_UPLOAD_SERVICE['timeout'])
            
            print(f"上传响应状态码: {response.status_code}")
            print(f"上传响应内容: {response.text}")
            
            db_file_path = None
            if response.status_code == 200:
                try:
                    upload_result = response.json()
                    print(f"上传结果: {upload_result}")
                    if upload_result.get('code') == 200 and upload_result.get('data'):
                        if isinstance(upload_result['data'], list) and upload_result['data']:
                            file_path = upload_result['data'][0]
                        elif isinstance(upload_result['data'], str):
                            file_path = upload_result['data']
                        print(f"获取到的文件路径: {file_path}")
                        
                        if file_path and 'upload/' in file_path:
                            db_file_path = file_path[file_path.find('upload/'):]
                        else:
                            db_file_path = file_path
                except Exception as json_error:
                    print(f"解析响应JSON失败: {json_error}")
            else:
                print(f"上传失败，状态码: {response.status_code}")
    except Exception as e:
        print(f"文件上传失败: {e}")
        import traceback
        traceback.print_exc()
        db_file_path = None
    
    # 无论上传是否成功，都继续生成报告
    print(f"文件路径: {file_path}")
    print(f"数据库存储路径: {db_file_path}")
    
    # 添加报告记录到数据库
    add_report(report_type, group_id if group_id != 'all' else None, group_name, report_name, db_file_path)
    
    # 重置文件指针
    output.seek(0)
    
    # 返回文件下载响应
    return send_file(
        output,
        download_name=filename,
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def generate_report_content(records, report_type, start_date_str=None):
    """生成报告内容"""
    import datetime
    
    # 解析开始时间
    if start_date_str:
        now = datetime.datetime.fromisoformat(start_date_str)
    else:
        now = datetime.datetime.now()
    
    # 根据报告类型确定时间范围
    if report_type == 'daily':
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = now.replace(hour=23, minute=59, second=59, microsecond=999999)
        title = f"服务器巡检日报 ({start_date.strftime('%Y-%m-%d')})"
    elif report_type == 'weekly':
        # 计算本周的开始和结束日期
        start_date = now - datetime.timedelta(days=now.weekday())
        start_date = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + datetime.timedelta(days=6, hours=23, minutes=59, seconds=59, microseconds=999999)
        title = f"服务器巡检周报 ({start_date.strftime('%Y-%m-%d')} 至 {end_date.strftime('%Y-%m-%d')})"
    elif report_type == 'monthly':
        # 计算本月的开始和结束日期
        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        if start_date.month == 12:
            end_date = start_date.replace(year=start_date.year + 1, month=1, day=1) - datetime.timedelta(seconds=1)
        else:
            end_date = start_date.replace(month=start_date.month + 1, day=1) - datetime.timedelta(seconds=1)
        title = f"服务器巡检月报 ({start_date.strftime('%Y-%m')})"
    else:
        return None
    
    # 筛选时间范围内的记录
    filtered_records = []
    for record in records:
        inspection_time = datetime.datetime.fromisoformat(record[7])
        if start_date <= inspection_time <= end_date:
            filtered_records.append(record)
    
    # 按服务器IP分组
    server_records = {}
    # 存储所有服务器IP，用于去重
    server_ips = set()
    for record in filtered_records:
        # 尝试获取服务器IP
        server_ip = None
        if len(record) > 10:
            for idx in range(9, min(len(record), 15)):
                if isinstance(record[idx], str):
                    import re
                    ip_pattern = r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$'
                    if re.match(ip_pattern, record[idx]):
                        server_ip = record[idx]
                        server_ips.add(server_ip)
                        break
        
        # 如果没有找到IP，使用服务器名称或其他标识符
        if not server_ip:
            server_ip = record[9] or record[10] or 'unknown'
            server_ips.add(server_ip)
        
        if server_ip not in server_records:
            server_records[server_ip] = []
        server_records[server_ip].append(record)
    
    return {
        'title': title,
        'start_date': start_date,
        'end_date': end_date,
        'server_records': server_records,
        'total_servers': len(server_ips),
        'total_inspections': len(filtered_records)
    }

def generate_report_docx(report_content, report_type):
    """生成报告Word文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(report_content['title'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加报告信息
    doc.add_paragraph()
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = '报告类型'
    info_table.cell(0, 1).text = {'daily': '日报', 'weekly': '周报', 'monthly': '月报'}.get(report_type, '未知')
    info_table.cell(1, 0).text = '开始时间'
    info_table.cell(1, 1).text = report_content['start_date'].strftime('%Y-%m-%d %H:%M:%S')
    info_table.cell(2, 0).text = '结束时间'
    info_table.cell(2, 1).text = report_content['end_date'].strftime('%Y-%m-%d %H:%M:%S')
    
    # 添加统计信息
    doc.add_paragraph()
    doc.add_heading('一、统计信息', level=2)
    doc.add_paragraph(f"巡检服务器数量: {report_content['total_servers']}")
    doc.add_paragraph(f"巡检次数: {report_content['total_inspections']}")
    
    # 添加服务器巡检详情
    doc.add_paragraph()
    doc.add_heading('二、服务器巡检详情', level=2)
    
    # 收集所有巡检记录
    all_records = []
    for server_name, records in report_content['server_records'].items():
        all_records.extend(records)
    
    # 添加巡检记录表格
    if all_records:
        table = doc.add_table(rows=len(all_records) + 1, cols=8)
        table.style = 'Table Grid'
        
        # 添加表头
        headers = ['序号', '服务器IP', '服务器名称', '巡检时间', 'CPU使用率', '内存使用率', '磁盘使用率', '状态']
        for col_idx, header in enumerate(headers):
            table.cell(0, col_idx).text = header
            table.cell(0, col_idx).paragraphs[0].runs[0].font.bold = True
        
        # 添加数据
        for row_idx, record in enumerate(all_records, 1):
            # 确保行索引不超出表格范围
            if row_idx < len(table.rows):
                table.cell(row_idx, 0).text = str(row_idx)  # 序号
                # 直接使用固定的索引位置获取服务器信息
                # 根据get_inspection_records函数的SQL查询，返回字段顺序为：
                # ir.*, s.name as server_name, s.ip as server_ip, g.name as group_name
                server_ip = '未知'
                server_name = '未知'
                
                # 检查记录长度，确保有足够的字段
                # 根据get_inspection_records函数的SQL查询，返回字段顺序为：
                # ir.*(12个字段) + s.name(服务器名称) + s.ip(服务器IP) + g.name(分组名称)
                if len(record) >= 15:
                    # 服务器名称在索引12，服务器IP在索引13
                    if len(record) > 12 and record[12]:
                        server_name = str(record[12])
                    if len(record) > 13 and record[13]:
                        server_ip = str(record[13])
                
                # 如果还是没有获取到服务器IP，尝试从记录中提取
                if server_ip == '未知':
                    # 尝试从描述字段中提取IP
                    if len(record) > 4 and isinstance(record[4], str):
                        import re
                        ip_pattern = r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'
                        match = re.search(ip_pattern, record[4])
                        if match:
                            server_ip = match.group(0)
                
                # 确保列索引不超出表格范围
                if len(table.columns) >= 8:
                    table.cell(row_idx, 1).text = server_ip  # 服务器IP
                    table.cell(row_idx, 2).text = server_name  # 服务器名称
                    table.cell(row_idx, 3).text = record[7]  # 巡检时间
                    table.cell(row_idx, 4).text = f"{record[4]}%" if record[4] else '无法获取'  # CPU使用率
                    
                    # 内存使用率
                    memory_usage = '无法获取'
                    if record[3]:
                        try:
                            mem_info = eval(record[3])
                            if mem_info.get('usage_percent'):
                                memory_usage = f"{mem_info['usage_percent']}%"
                        except:
                            pass
                    table.cell(row_idx, 5).text = memory_usage
                    
                    # 磁盘使用率
                    disk_usage = '无法获取'
                    if record[2]:
                        try:
                            disk_info = eval(record[2])
                            if disk_info:
                                max_usage = max([d.get('usage_percent', 0) for d in disk_info])
                                disk_usage = f"{max_usage}%"
                        except:
                            pass
                    table.cell(row_idx, 6).text = disk_usage
                    
                    # 状态
                    status = '正常'
                    if record[8] and record[8] != '无告警':
                        status = '告警'
                    table.cell(row_idx, 7).text = status
                    
                    # 高亮告警行
                    if status == '告警':
                        for col_idx in range(8):
                            table.cell(row_idx, col_idx).paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
        
    # 添加总结
    doc.add_paragraph()
    doc.add_heading('三、总结', level=2)
    
    # 统计告警次数
    alert_count = 0
    for records in report_content['server_records'].values():
        for record in records:
            if record[8] and record[8] != '无告警':
                alert_count += 1
    
    if alert_count == 0:
        doc.add_paragraph('所有服务器运行正常，未发现告警。')
    else:
        doc.add_paragraph(f'共发现 {alert_count} 次告警，需要关注。')
    
    # 添加页脚
    footer = doc.sections[0].footer
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = "服务器自动巡检系统生成"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def get_connection():
    import sqlite3
    from config import DATABASE_PATH
    return sqlite3.connect(DATABASE_PATH)

# ========== 故障处置相关路由 ==========
@app.route('/faults')
@login_required
def faults():
    groups = get_all_groups()
    
    # 获取筛选参数
    group_id = request.args.get('group_id')
    fault_type = request.args.get('fault_type')
    fault_level = request.args.get('fault_level')
    start_time = request.args.get('start_time')
    end_time = request.args.get('end_time')
    keyword = request.args.get('keyword')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    # 获取故障列表
    faults_list = get_all_faults(page, per_page, group_id, fault_type, fault_level, start_time, end_time, keyword)
    total = get_faults_count(group_id, fault_type, fault_level, start_time, end_time, keyword)
    
    return render_template('faults.html', 
                         groups=groups, 
                         faults=faults_list, 
                         page=page, 
                         per_page=per_page, 
                         total=total)

@app.route('/api/add_fault', methods=['POST'])
@login_required
def api_add_fault():
    data = request.json
    
    group_id = data.get('group_id')
    group_name = data.get('group_name')
    event_name = data.get('event_name')
    fault_level = data.get('fault_level')
    fault_type = data.get('fault_type')
    fault_description = data.get('fault_description')
    occurrence_time = data.get('occurrence_time')
    discoverer = data.get('discoverer')
    
    if not event_name or not fault_description or not occurrence_time or not discoverer:
        return jsonify({'success': False, 'message': '必填字段不能为空'})
    
    fault_id = add_fault(group_id, group_name, event_name, fault_level, fault_type, fault_description, occurrence_time, discoverer)
    if fault_id:
        return jsonify({'success': True, 'message': '添加成功'})
    else:
        return jsonify({'success': False, 'message': '添加失败'})

@app.route('/api/get_fault/<int:fault_id>')
@login_required
def api_get_fault(fault_id):
    fault = get_fault_by_id(fault_id)
    if fault:
        return jsonify({'success': True, 'data': fault})
    else:
        return jsonify({'success': False, 'message': '记录不存在'})

@app.route('/api/update_fault', methods=['POST'])
@login_required
def api_update_fault():
    data = request.json
    
    fault_id = data.get('id')
    if not fault_id:
        return jsonify({'success': False, 'message': 'ID不能为空'})
    
    update_data = {}
    if 'group_id' in data:
        update_data['group_id'] = data['group_id']
    if 'group_name' in data:
        update_data['group_name'] = data['group_name']
    if 'event_name' in data:
        update_data['event_name'] = data['event_name']
    if 'fault_level' in data:
        update_data['fault_level'] = data['fault_level']
    if 'fault_type' in data:
        update_data['fault_type'] = data['fault_type']
    if 'fault_description' in data:
        update_data['fault_description'] = data['fault_description']
    if 'occurrence_time' in data:
        update_data['occurrence_time'] = data['occurrence_time']
    if 'discoverer' in data:
        update_data['discoverer'] = data['discoverer']
    if 'processing_status' in data:
        update_data['processing_status'] = data['processing_status']
    if 'processing_person' in data:
        update_data['processing_person'] = data['processing_person']
    if 'processing_time' in data:
        update_data['processing_time'] = data['processing_time']
    if 'processing_description' in data:
        update_data['processing_description'] = data['processing_description']
    if 'report_path' in data:
        update_data['report_path'] = data['report_path']
    
    result = update_fault(fault_id, **update_data)
    if result:
        return jsonify({'success': True, 'message': '更新成功'})
    else:
        return jsonify({'success': False, 'message': '更新失败'})

@app.route('/api/delete_fault/<int:fault_id>')
@login_required
def api_delete_fault(fault_id):
    result = delete_fault(fault_id)
    if result:
        return jsonify({'success': True, 'message': '删除成功'})
    else:
        return jsonify({'success': False, 'message': '删除失败'})

def generate_fault_report(fault_data):
    """生成故障处置报告"""
    doc = Document()
    
    title = doc.add_heading('故障处置报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    
    def set_label_and_content(cell, label, content=''):
        if not cell.paragraphs:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        run_label = p.add_run(label)
        run_label.bold = True
        if content:
            p.add_run('\n' + content)
    
    set_label_and_content(table.cell(0, 0), '事件名称')
    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(0, 1).text = fault_data.get('event_name') or ''
    
    set_label_and_content(table.cell(1, 0), '事件时间')
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 1).text = fault_data.get('occurrence_time') or ''
    
    row2 = table.rows[2]
    row2.cells[0].merge(row2.cells[3])
    set_label_and_content(row2.cells[0], '故障事件描述：', fault_data.get('fault_description') or '')
    
    row3 = table.rows[3]
    row3.cells[0].merge(row3.cells[3])
    set_label_and_content(row3.cells[0], '发生原因分析：', fault_data.get('processing_description') or '')
    
    row4 = table.rows[4]
    row4.cells[0].merge(row4.cells[3])
    set_label_and_content(row4.cells[0], '影响范围：')
    
    row5 = table.rows[5]
    row5.cells[0].merge(row5.cells[3])
    set_label_and_content(row5.cells[0], '处置措施和过程：')
    
    row6 = table.rows[6]
    row6.cells[0].merge(row6.cells[3])
    set_label_and_content(row6.cells[0], '最终结果：', fault_data.get('processing_status') or '')
    
    row7 = table.rows[7]
    row7.cells[0].merge(row7.cells[3])
    set_label_and_content(row7.cells[0], '后续措施：')
    
    row8 = table.rows[8]
    row8.cells[0].merge(row8.cells[3])
    set_label_and_content(row8.cells[0], '相关建议：')
    
    set_label_and_content(table.cell(9, 0), '故障处理人员')
    table.cell(9, 1).text = fault_data.get('processing_person') or ''
    set_label_and_content(table.cell(9, 2), '审核人员')
    table.cell(9, 3).text = fault_data.get('review_person') or ''
    
    set_label_and_content(table.cell(10, 0), '审核意见')
    table.cell(10, 1).text = fault_data.get('review_comment') or ''
    set_label_and_content(table.cell(10, 2), '审核时间')
    table.cell(10, 3).text = datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
    
    for i in range(2, 9):
        table.rows[i].height = Pt(60)
    
    return doc

@app.route('/api/review_fault', methods=['POST'])
@login_required
def api_review_fault():
    data = request.json
    
    fault_id = data.get('fault_id')
    review_person = data.get('review_person')
    review_comment = data.get('review_comment')
    
    if not fault_id or not review_person or not review_comment:
        return jsonify({'success': False, 'message': '参数不能为空'})
    
    fault = get_fault_by_id(fault_id)
    if not fault:
        return jsonify({'success': False, 'message': '故障记录不存在'})
    
    fault_dict = {
        'id': fault[0],
        'group_id': fault[1],
        'group_name': fault[2],
        'event_name': fault[3],
        'fault_level': fault[4],
        'fault_type': fault[5],
        'fault_description': fault[6],
        'occurrence_time': fault[7],
        'discoverer': fault[8],
        'processing_status': fault[9],
        'processing_person': fault[10],
        'processing_time': fault[11],
        'processing_description': fault[12],
        'review_status': '已审核',
        'review_person': review_person,
        'review_comment': review_comment
    }
    
    doc = generate_fault_report(fault_dict)
    
    now = datetime.datetime.now()
    year_month = now.strftime('%Y-%m')
    upload_dir = os.path.join(config.UPLOAD_LOCAL_PATH, year_month)
    os.makedirs(upload_dir, exist_ok=True)
    
    filename = f"故障处置报告_{fault_id}_{now.strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(upload_dir, filename)
    
    doc.save(file_path)
    
    relative_path = f"upload/{year_month}/{filename}"
    
    result = review_fault(fault_id, review_person, review_comment, relative_path)
    if result:
        return jsonify({'success': True, 'message': '审核成功，报告已生成'})
    else:
        return jsonify({'success': False, 'message': '审核失败'})

# ==================== 短信告警功能 ====================

def send_inspection_alert_sms(server, result):
    """
    发送巡检告警短信
    :param server: 服务器信息
    :param result: 巡检结果
    """
    from sms_notifier import send_alert_sms
    
    # 获取告警配置
    config = get_alert_config()
    if not config or not config[2]:  # is_enabled
        return
    
    # 构建配置字典
    sms_config = {
        'provider': config[1],
        'sign_name': config[7],
        'inspection_template_code': config[9],
        'phone_numbers': config[12]
    }
    
    if config[1] == 'aliyun':
        sms_config['access_key_id'] = config[3]
        sms_config['access_key_secret'] = config[4]
    elif config[1] == 'tencent':
        sms_config['secret_id'] = config[5]
        sms_config['secret_key'] = config[6]
        sms_config['app_id'] = config[8]
    elif config[1] == 'baidu':
        sms_config['access_key_id'] = config[3]
        sms_config['secret_access_key'] = config[4]
    
    # 解析告警内容
    alert_content = result.get('alert_content', '')
    alert_items = []
    
    # 提取告警项
    if 'CPU' in alert_content:
        cpu_match = alert_content.split('CPU使用率超过阈值')[0] if 'CPU使用率超过阈值' in alert_content else ''
        if cpu_match:
            alert_items.append('CPU使用率')
    
    if '内存' in alert_content:
        alert_items.append('内存使用率')
    
    if '磁盘' in alert_content:
        alert_items.append('磁盘使用率')
    
    alert_item_str = '、'.join(alert_items) if alert_items else '系统资源'
    
    # 获取告警值
    alert_value = '超过阈值'
    if result.get('cpu_usage'):
        try:
            cpu_val = float(result['cpu_usage'])
            if cpu_val > config.MAX_ALERT_THRESHOLD:
                alert_value = f"CPU {cpu_val}%"
        except:
            pass
    
    # 构建告警数据
    alert_data = {
        'server_name': server[1],
        'server_ip': server[2],
        'alert_item': alert_item_str,
        'alert_value': alert_value,
        'threshold': str(config.MAX_ALERT_THRESHOLD),
        'time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    }
    
    # 发送短信
    success, message = send_alert_sms(
        phone_numbers=config[12],
        alert_type='inspection',
        alert_data=alert_data,
        config=sms_config
    )
    
    # 记录发送日志
    add_alert_log(
        alert_type='inspection',
        alert_title=f"服务器巡检告警 - {server[1]}",
        alert_content=alert_content,
        phone_numbers=config[12],
        send_status='成功' if success else '失败',
        send_result=message
    )

def send_fault_alert_sms(fault_data):
    """
    发送故障告警短信
    :param fault_data: 故障数据字典
    """
    from sms_notifier import send_alert_sms
    
    # 获取告警配置
    config = get_alert_config()
    if not config or not config[2]:  # is_enabled
        return
    
    # 构建配置字典
    sms_config = {
        'provider': config[1],
        'sign_name': config[7],
        'fault_template_code': config[10],
        'phone_numbers': config[12]
    }
    
    if config[1] == 'aliyun':
        sms_config['access_key_id'] = config[3]
        sms_config['access_key_secret'] = config[4]
    elif config[1] == 'tencent':
        sms_config['secret_id'] = config[5]
        sms_config['secret_key'] = config[6]
        sms_config['app_id'] = config[8]
    elif config[1] == 'baidu':
        sms_config['access_key_id'] = config[3]
        sms_config['secret_access_key'] = config[4]
    
    # 发送短信
    success, message = send_alert_sms(
        phone_numbers=config[12],
        alert_type='fault',
        alert_data=fault_data,
        config=sms_config
    )
    
    # 记录发送日志
    add_alert_log(
        alert_type='fault',
        alert_title=f"故障告警 - {fault_data.get('fault_name', '')}",
        alert_content=fault_data.get('description', ''),
        phone_numbers=config[12],
        send_status='成功' if success else '失败',
        send_result=message
    )

# ==================== 告警通知配置路由 ====================

@app.route('/alert_config')
@login_required
def alert_config():
    """告警通知配置页面"""
    config = get_alert_config()
    return render_template('alert_config.html', config=config)

@app.route('/alert_config', methods=['POST'])
@login_required
def alert_config_save():
    """保存告警通知配置"""
    try:
        provider = request.form.get('provider', '').strip()
        is_enabled = 1 if request.form.get('is_enabled') else 0
        phone_numbers = request.form.get('phone_numbers', '').strip()
        
        # 根据服务商获取对应的配置字段
        access_key_id = ''
        access_key_secret = ''
        secret_id = ''
        secret_key = ''
        sign_name = ''
        app_id = ''
        
        if provider == 'aliyun':
            access_key_id = request.form.get('access_key_id', '').strip()
            access_key_secret = request.form.get('access_key_secret', '').strip()
            sign_name = request.form.get('sign_name', '').strip()
        elif provider == 'tencent':
            secret_id = request.form.get('secret_id', '').strip()
            secret_key = request.form.get('secret_key', '').strip()
            sign_name = request.form.get('sign_name_tencent', '').strip()
            app_id = request.form.get('app_id', '').strip()
        elif provider == 'baidu':
            access_key_id = request.form.get('access_key_id_baidu', '').strip()
            access_key_secret = request.form.get('secret_access_key', '').strip()
            sign_name = request.form.get('sign_name_baidu', '').strip()
        elif provider == 'custom':
            sign_name = request.form.get('sign_name', '').strip()
        
        inspection_template_code = request.form.get('inspection_template_code', '').strip()
        fault_template_code = request.form.get('fault_template_code', '').strip()
        system_template_code = request.form.get('system_template_code', '').strip()
        
        custom_api_url = request.form.get('custom_api_url', '').strip()
        custom_headers = request.form.get('custom_headers', '').strip()
        
        save_alert_config(
            provider=provider,
            is_enabled=is_enabled,
            access_key_id=access_key_id,
            access_key_secret=access_key_secret,
            secret_id=secret_id,
            secret_key=secret_key,
            sign_name=sign_name,
            app_id=app_id,
            inspection_template_code=inspection_template_code,
            fault_template_code=fault_template_code,
            system_template_code=system_template_code,
            phone_numbers=phone_numbers,
            custom_api_url=custom_api_url,
            custom_headers=custom_headers
        )
        
        config = get_alert_config()
        return render_template('alert_config.html', config=config, message='配置保存成功', success=True)
    except Exception as e:
        config = get_alert_config()
        return render_template('alert_config.html', config=config, message=f'保存失败: {str(e)}', success=False)

@app.route('/api/test_sms', methods=['POST'])
@login_required
def api_test_sms():
    """测试短信发送"""
    try:
        from sms_notifier import get_sms_notifier, send_alert_sms
        
        provider = request.form.get('provider', '').strip()
        phone_numbers = request.form.get('phone_numbers', '').strip()
        
        if not provider or not phone_numbers:
            return jsonify({'success': False, 'message': '请填写完整配置信息'})
        
        # 构建配置字典
        config = {
            'provider': provider,
            'sign_name': request.form.get('sign_name') or request.form.get('sign_name_tencent') or request.form.get('sign_name_baidu', '').strip()
        }
        
        if provider == 'aliyun':
            config['access_key_id'] = request.form.get('access_key_id', '').strip()
            config['access_key_secret'] = request.form.get('access_key_secret', '').strip()
        elif provider == 'tencent':
            config['secret_id'] = request.form.get('secret_id', '').strip()
            config['secret_key'] = request.form.get('secret_key', '').strip()
            config['app_id'] = request.form.get('app_id', '').strip()
        elif provider == 'baidu':
            config['access_key_id'] = request.form.get('access_key_id_baidu', '').strip()
            config['secret_access_key'] = request.form.get('secret_access_key', '').strip()
        elif provider == 'custom':
            config['custom_api_url'] = request.form.get('custom_api_url', '').strip()
            config['custom_headers'] = request.form.get('custom_headers', '').strip()
        
        # 使用系统告警模板发送测试短信
        config['system_template_code'] = request.form.get('system_template_code', '').strip()
        
        if not config['system_template_code']:
            return jsonify({'success': False, 'message': '请先配置系统告警模板代码'})
        
        # 发送测试短信
        alert_data = {
            'alert_title': '测试告警',
            'alert_content': '这是一条测试短信，用于验证告警通知配置是否正确。',
            'time': datetime.now().strftime('%Y-%m-%d %H:%M')
        }
        
        success, message = send_alert_sms(phone_numbers, 'system', alert_data, config)
        
        # 记录发送日志
        add_alert_log(
            alert_type='system',
            alert_title='测试短信',
            alert_content=alert_data['alert_content'],
            phone_numbers=phone_numbers,
            send_status='成功' if success else '失败',
            send_result=message
        )
        
        return jsonify({'success': success, 'message': message})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'测试发送失败: {str(e)}'})

@app.route('/alert_logs')
@login_required
def alert_logs():
    """告警记录页面"""
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    logs = get_alert_logs(page=page, per_page=per_page)
    total = get_alert_logs_count()
    
    return render_template('alert_logs.html', logs=logs, page=page, per_page=per_page, total=total)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=config.PORT)